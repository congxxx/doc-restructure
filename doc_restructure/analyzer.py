"""Document structure analyzer."""

from docx import Document
from .utils import (
    detect_style_mapping,
    get_heading_level,
    get_elem_tag,
    get_body_elements,
)


def analyze(doc_path):
    """Analyze document structure and print summary.

    Args:
        doc_path: Path to .docx file
    """
    doc = Document(doc_path)
    style_mapping = detect_style_mapping(doc)
    paragraphs = doc.paragraphs
    body_elements = get_body_elements(doc)

    chapter_starts = [
        i for i, p in enumerate(paragraphs)
        if p.style and p.style.name == "Heading 1"
    ]

    print(f"Document: {doc_path}")
    print(f"Paragraphs: {len(paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Chapters: {len(chapter_starts)}")

    # Check for missing tables (table caption without following table)
    missing_tables = _check_missing_tables(body_elements)
    if missing_tables:
        print(f"\nWARNING: {len(missing_tables)} table captions without tables!")
        for caption in missing_tables[:5]:
            print(f"  - {caption}")

    # Print heading structure
    print("\n=== Heading Structure ===")
    for ci in range(len(chapter_starts)):
        s = chapter_starts[ci]
        e = chapter_starts[ci + 1] if ci + 1 < len(chapter_starts) else len(paragraphs)

        figs = sum(
            1 for i in range(s, e)
            if paragraphs[i].style and "图注" in paragraphs[i].style.name
            and paragraphs[i].text.strip()
        )
        tabs = sum(
            1 for i in range(s, e)
            if paragraphs[i].style and "表注" in paragraphs[i].style.name
            and paragraphs[i].text.strip()
        )

        ch_title = paragraphs[s].text.strip()[:50]
        print(f"\n{ch_title} ({figs} figs, {tabs} tables)")

        for i in range(s, min(e, s + 100)):
            p = paragraphs[i]
            level = get_heading_level(p)
            if level > 0:
                indent = "  " * level
                print(f"{indent}[H{level}] {p.text.strip()[:60]}")

    # Print figure/table list
    print("\n=== Figures & Tables ===")
    for ci in range(len(chapter_starts)):
        s = chapter_starts[ci]
        e = chapter_starts[ci + 1] if ci + 1 < len(chapter_starts) else len(paragraphs)

        for i in range(s, e):
            p = paragraphs[i]
            if p.style and ("图注" in p.style.name or "表注" in p.style.name):
                if p.text.strip():
                    print(f"  [{i}] {p.text.strip()[:60]}")


def _check_missing_tables(body_elements):
    """Check for table captions not followed by tables.

    Returns:
        List of caption texts that have no following table
    """
    from docx.oxml.ns import qn
    missing = []
    for i, elem in enumerate(body_elements):
        tag = get_elem_tag(elem)
        if tag == "p":
            # Check if this is a table caption (style ID 178)
            pPr = elem.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None and pStyle.get(qn("w:val")) == "178":
                    # Check if next element is a table
                    if i + 1 >= len(body_elements) or get_elem_tag(body_elements[i + 1]) != "tbl":
                        # Get caption text
                        texts = []
                        from docx.oxml.ns import qn
                        for r in elem.findall(qn("w:r")):
                            for t in r.findall(qn("w:t")):
                                if t.text:
                                    texts.append(t.text)
                        missing.append("".join(texts)[:50])
    return missing
