"""Shared utilities for document manipulation."""

import copy
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Cache for style ID mapping per document
_style_cache = {}


def detect_style_mapping(doc):
    """Detect heading level to style ID mapping from a document.

    Word/WPS uses numeric style IDs internally (e.g., Heading 1 = "3").
    This function reads the actual mapping from the document's styles.

    Returns:
        dict: {level: style_id_string}, e.g., {1: "3", 2: "4", 3: "5"}
    """
    doc_id = id(doc)
    if doc_id in _style_cache:
        return _style_cache[doc_id]

    mapping = {}
    for style in doc.styles:
        if "Heading" in style.name:
            try:
                level = int(style.name.split()[-1])
                mapping[level] = str(style.style_id)
            except (ValueError, AttributeError):
                pass

    _style_cache[doc_id] = mapping
    return mapping


def get_heading_level(para):
    """Get heading level from a paragraph (1-9, 0 if not a heading)."""
    if not para.style:
        return 0
    if "Heading" in para.style.name:
        try:
            return int(para.style.name.split()[-1])
        except ValueError:
            return 0
    return 0


def get_level_from_elem(elem, style_mapping):
    """Get heading level from a paragraph XML element."""
    pPr = elem.find(qn("w:pPr"))
    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            sid = pStyle.get(qn("w:val"), "")
            for lvl, lid in style_mapping.items():
                if lid == sid:
                    return lvl
    return 0


def set_heading_level(elem, target_level, style_mapping):
    """Set heading level on a paragraph XML element."""
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return
    style_id = style_mapping.get(target_level, str(target_level))
    pStyle.set(qn("w:val"), style_id)


def make_heading(text, level, template_elem, style_mapping):
    """Create a new heading element based on a template.

    Args:
        text: Heading text
        level: Heading level (1-9)
        template_elem: XML element to use as formatting template
        style_mapping: Style ID mapping from detect_style_mapping()

    Returns:
        New XML element with the specified heading text and level
    """
    new_elem = copy.deepcopy(template_elem)
    set_heading_level(new_elem, level, style_mapping)

    # Clear all existing text
    for r in new_elem.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            t.text = ""

    # Set new text on the last run
    r_elems = new_elem.findall(qn("w:r"))
    if r_elems:
        last_r = r_elems[-1]
        t_elems = last_r.findall(qn("w:t"))
        if t_elems:
            t_elems[0].text = text
        else:
            t = OxmlElement("w:t")
            t.text = text
            last_r.append(t)

    return new_elem


def replace_across_runs(para, old_text, new_text):
    """Replace text that may span multiple runs.

    Word often splits text across multiple runs (e.g., "图1-2" might be
    split as "图1-" in run[0] and "2" in run[1]). This function handles
    that by first trying single-run replacement, then falling back to
    full-text replacement.

    Args:
        para: Paragraph object
        old_text: Text to find
        new_text: Replacement text

    Returns:
        True if replacement was made
    """
    # Try single-run replacement first (preserves formatting)
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # Fallback: merge all runs, replace, put back in first run
    # Note: this may lose per-run formatting
    full = "".join(r.text for r in para.runs)
    if old_text in full:
        new_full = full.replace(old_text, new_text)
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""
        return True

    return False


def get_elem_tag(elem):
    """Get the local tag name from an XML element (without namespace)."""
    return elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag


def get_body_elements(doc):
    """Get all body elements (paragraphs, tables, etc.) in document order.

    Unlike doc.paragraphs which only returns paragraph elements,
    this returns ALL body elements including tables.
    """
    return list(doc.element.body)


def get_para_to_body_map(doc):
    """Build mapping from paragraph index to body element index."""
    para_elements = [p._element for p in doc.paragraphs]
    body_elements = get_body_elements(doc)
    return {id(e): i for i, e in enumerate(body_elements)}


def get_body_to_para_map(doc):
    """Build mapping from body element ID to paragraph index."""
    para_elements = [p._element for p in doc.paragraphs]
    return {id(pe): i for i, pe in enumerate(para_elements)}
