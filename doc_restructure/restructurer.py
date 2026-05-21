"""Document structure reorganizer.

Moves sections from source document to a new structure defined by a mapping.
Preserves all content (paragraphs, tables, images) and formatting.
"""

import copy
from docx import Document
from docx.oxml.ns import qn

from .utils import (
    detect_style_mapping,
    get_heading_level,
    get_level_from_elem,
    set_heading_level,
    make_heading,
    get_elem_tag,
    get_body_elements,
)


def restructure(source_path, output_path, mapping):
    """Restructure a document according to a mapping definition.

    Args:
        source_path: Path to source .docx file
        output_path: Path to output .docx file
        mapping: Dict with 'chapters' key, each containing 'title' and 'sections'
            See README.md for mapping format details.
    """
    doc = Document(source_path)
    style_mapping = detect_style_mapping(doc)
    paragraphs = doc.paragraphs
    body_elements = get_body_elements(doc)
    para_elements = [p._element for p in paragraphs]

    # Build index mappings
    elem_to_body_idx = {id(e): i for i, e in enumerate(body_elements)}
    body_elem_to_para_idx = {id(pe): i for i, pe in enumerate(para_elements)}

    # Find chapter boundaries
    chapter_starts = [
        i for i, p in enumerate(paragraphs)
        if p.style and p.style.name == "Heading 1"
    ]

    # Template for new headings
    h1_template = para_elements[chapter_starts[0]]

    # Create new document from template (preserves styles)
    new_doc = Document(source_path)
    new_body = new_doc.element.body

    # Clear existing content
    for child in list(new_body):
        new_body.remove(child)

    # Track collected paragraphs to avoid duplicates
    collected = {}

    def get_chapter_range(chap_idx):
        """Get paragraph range [start, end) for a chapter."""
        s = chapter_starts[chap_idx]
        e = (
            chapter_starts[chap_idx + 1]
            if chap_idx + 1 < len(chapter_starts)
            else len(paragraphs)
        )
        return s, e

    def find_heading(chap_idx, text):
        """Find a heading by text within a chapter (partial match).

        Returns:
            (para_idx, already_collected) tuple, or (None, False) if not found.
        """
        s, e = get_chapter_range(chap_idx)
        for i in range(s, e):
            p = paragraphs[i]
            if p.style and "Heading" in p.style.name and text in p.text:
                is_collected = chap_idx in collected and i in collected[chap_idx]
                return i, is_collected
        return None, False

    def collect_section(chap_idx, heading_text, target_level):
        """Collect all body elements under a heading.

        Includes:
        - Intro paragraphs (between H1 and first H2) if this is the first
          section collected from the chapter
        - The heading itself
        - All content until the next same-level or higher-level heading
        - Trailing tables after figure/table captions

        Args:
            chap_idx: Source chapter index
            heading_text: Text to match in headings
            target_level: Target heading level in new document

        Returns:
            List of XML elements (deep copies) ready to append to new body.
            Returns None if heading was already collected (silent skip).
        """
        start_idx, already_collected = find_heading(chap_idx, heading_text)
        if start_idx is None:
            return []
        if already_collected:
            return None  # Signal: already included by parent heading

        start_level = get_heading_level(paragraphs[start_idx])
        s, e = get_chapter_range(chap_idx)

        if chap_idx not in collected:
            collected[chap_idx] = set()

        para_idx_list = []

        # Collect intro paragraphs (between H1 and first H2)
        if s not in collected[chap_idx]:
            collected[chap_idx].add(s)
            for i in range(s + 1, e):
                p = paragraphs[i]
                if p.style and "Heading" in p.style.name:
                    break
                if i not in collected[chap_idx]:
                    para_idx_list.append(i)
                    collected[chap_idx].add(i)

        # Add the heading itself
        if start_idx not in collected[chap_idx]:
            para_idx_list.append(start_idx)
            collected[chap_idx].add(start_idx)

        # Collect content until next same-level or higher-level heading
        for i in range(start_idx + 1, e):
            p = paragraphs[i]
            if p.style and "Heading" in p.style.name:
                if get_heading_level(p) <= start_level:
                    break
            if i not in collected[chap_idx]:
                para_idx_list.append(i)
                collected[chap_idx].add(i)

        if not para_idx_list:
            return []

        # Map paragraph indices to body element indices
        first_elem = para_elements[para_idx_list[0]]
        last_elem = para_elements[para_idx_list[-1]]
        first_bi = elem_to_body_idx.get(id(first_elem))
        last_bi = elem_to_body_idx.get(id(last_elem))

        # KEY FIX: Extend range to include trailing tables after captions
        # Tables are separate body elements not in the paragraphs list.
        # If the last paragraph is a figure/table caption, include the
        # following table/image element.
        if last_bi is not None:
            last_para = paragraphs[para_idx_list[-1]]
            if last_para.style and (
                "表注" in last_para.style.name or "图注" in last_para.style.name
            ):
                while last_bi + 1 < len(body_elements):
                    next_tag = get_elem_tag(body_elements[last_bi + 1])
                    if next_tag == "tbl":
                        last_bi += 1
                    else:
                        break

        if first_bi is None or last_bi is None:
            return []

        # Calculate level offset for heading adjustment
        level_offset = target_level - start_level

        # Collect and transform body elements
        result = []
        for bi in range(first_bi, last_bi + 1):
            elem = body_elements[bi]
            tag = get_elem_tag(elem)
            new_elem = copy.deepcopy(elem)

            if tag == "p":
                orig_level = get_level_from_elem(new_elem, style_mapping)

                # Skip parent H2 headings (the new section H2 replaces them)
                if orig_level == 2:
                    pi = body_elem_to_para_idx.get(id(elem))
                    if pi is not None and pi not in para_idx_list:
                        continue

                # Adjust heading level
                if orig_level > 0:
                    new_level = orig_level + level_offset
                    set_heading_level(new_elem, max(1, min(9, new_level)), style_mapping)

            result.append(new_elem)

        return result

    # Execute restructuring
    for chapter in mapping["chapters"]:
        h1 = make_heading(chapter["title"], 1, h1_template, style_mapping)
        new_body.append(h1)
        print(f"\n=== {chapter['title']} ===")

        chap_collected = {}
        for section in chapter["sections"]:
            h2 = make_heading(section["title"], 2, h1_template, style_mapping)
            new_body.append(h2)

            for source in section["sources"]:
                ci = source["chapter"]
                ht = source["heading"]
                tl = source.get("target_level", 3)

                if ci not in chap_collected:
                    chap_collected[ci] = set()
                collected[ci] = chap_collected[ci]

                elements = collect_section(ci, ht, tl)
                if elements is None:
                    pass  # Already collected by parent heading, silently skip
                elif elements:
                    print(f"  {ht} ({len(elements)} elements)")
                    for elem in elements:
                        new_body.append(elem)
                else:
                    print(f"  WARNING: not found: {ht}")

    new_doc.save(output_path)
    print(f"\nSaved: {output_path}")
