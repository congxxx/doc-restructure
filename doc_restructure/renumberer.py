"""Figure/table and heading renumbering."""

import re
from docx import Document
from .utils import detect_style_mapping, get_heading_level


# Regex patterns for figure and table IDs
FIG_RE = re.compile(r"(图)\s*(\d+)\s*-\s*(\d+)")
TAB_RE = re.compile(r"(表)\s*(\d+)\s*-\s*(\d+)")


def renumber_figures(doc_path):
    """Renumber figures and tables sequentially within each chapter.

    Figures are numbered as: 图{chapter}-{sequence}
    Tables are numbered as: 表{chapter}-{sequence}

    Also updates all cross-references in body text.

    Args:
        doc_path: Path to .docx file (modified in place)
    """
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    chapter_starts = [
        i for i, p in enumerate(paragraphs)
        if p.style and p.style.name == "Heading 1"
    ]

    # Step 1: Build caption changes and per-chapter remap lists
    caption_changes = []
    chapter_remaps = {}

    for ci in range(len(chapter_starts)):
        s = chapter_starts[ci]
        e = chapter_starts[ci + 1] if ci + 1 < len(chapter_starts) else len(paragraphs)
        new_ch = ci + 1
        fig_counter = 0
        tab_counter = 0
        remaps = []

        for i in range(s, e):
            p = paragraphs[i]
            if not p.style:
                continue
            text = p.text.strip()
            if not text:
                continue

            if "图注" in p.style.name:
                m = FIG_RE.search(text)
                if m:
                    old_id = f"图{m.group(2)}-{m.group(3)}"
                    fig_counter += 1
                    new_id = f"图{new_ch}-{fig_counter}"
                    if old_id != new_id:
                        caption_changes.append((i, old_id, new_id))
                        remaps.append((old_id, new_id))

            elif "表注" in p.style.name:
                m = TAB_RE.search(text)
                if m:
                    old_id = f"表{m.group(2)}-{m.group(3)}"
                    tab_counter += 1
                    new_id = f"表{new_ch}-{tab_counter}"
                    if old_id != new_id:
                        caption_changes.append((i, old_id, new_id))
                        remaps.append((old_id, new_id))

        chapter_remaps[ci] = remaps
        print(f"Chapter {new_ch}: {fig_counter} figures, {tab_counter} tables")

    # Step 2: Replace caption text (handles cross-run replacement)
    print(f"\nReplacing {len(caption_changes)} captions...")
    for para_idx, old_id, new_id in caption_changes:
        p = paragraphs[para_idx]
        old_type = old_id[0]
        old_num = old_id[1:]
        pattern = re.escape(old_type) + r"\s*" + re.escape(old_num)
        full_text = "".join(r.text for r in p.runs)
        new_full = re.sub(pattern, new_id, full_text)
        if new_full != full_text and p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""

    # Step 3: Replace body text references
    print("Replacing body text references...")
    total_changes = 0

    for ci in range(len(chapter_starts)):
        s = chapter_starts[ci]
        e = chapter_starts[ci + 1] if ci + 1 < len(chapter_starts) else len(paragraphs)
        remaps = chapter_remaps[ci]
        if not remaps:
            continue

        # Group remaps by old_id (handles duplicates)
        id_seqs = {}
        for old_id, new_id in remaps:
            id_seqs.setdefault(old_id, []).append(new_id)

        sorted_ids = sorted(id_seqs.keys(), key=len, reverse=True)
        changes = 0

        for i in range(s, e):
            p = paragraphs[i]
            # Skip captions (already processed)
            if p.style and ("图注" in p.style.name or "表注" in p.style.name):
                continue

            text = p.text
            if not text:
                continue

            # Build replacement
            new_text = text
            for old_id in sorted_ids:
                new_id = id_seqs[old_id][0]
                old_type = old_id[0]
                old_num = old_id[1:]
                pattern = re.escape(old_type) + r"\s*" + re.escape(old_num)
                new_text = re.sub(pattern, new_id, new_text)

            if new_text != text:
                changes += 1
                # Apply to runs
                if p.runs:
                    run_text = "".join(r.text for r in p.runs)
                    for old_id in sorted_ids:
                        new_id = id_seqs[old_id][0]
                        old_type = old_id[0]
                        old_num = old_id[1:]
                        pattern = re.escape(old_type) + r"\s*" + re.escape(old_num)
                        run_text = re.sub(pattern, new_id, run_text)

                    # Try single-run replacement first
                    any_match = False
                    for run in p.runs:
                        run_new = run.text
                        for old_id in sorted_ids:
                            new_id = id_seqs[old_id][0]
                            old_type = old_id[0]
                            old_num = old_id[1:]
                            pattern = re.escape(old_type) + r"\s*" + re.escape(old_num)
                            run_new = re.sub(pattern, new_id, run_new)
                        if run_new != run.text:
                            run.text = run_new
                            any_match = True

                    # Fallback: full text replacement
                    if not any_match and p.runs:
                        p.runs[0].text = run_text
                        for r in p.runs[1:]:
                            r.text = ""

        total_changes += changes

    print(f"Modified {total_changes} paragraphs")
    doc.save(doc_path)
    print(f"Saved: {doc_path}")


def renumber_headings(doc_path):
    """Renumber headings hierarchically within each chapter.

    H2: {chapter}.{h2_counter}
    H3: {chapter}.{h2_counter}.{h3_counter}
    H4: {chapter}.{h2_counter}.{h3_counter}.{h4_counter}

    Counters reset at each parent level.

    Args:
        doc_path: Path to .docx file (modified in place)
    """
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    chapter_starts = [
        i for i, p in enumerate(paragraphs)
        if p.style and p.style.name == "Heading 1"
    ]

    changes = 0
    for ci in range(len(chapter_starts)):
        s = chapter_starts[ci]
        e = chapter_starts[ci + 1] if ci + 1 < len(chapter_starts) else len(paragraphs)
        ch_num = ci + 1

        h2_counter = 0
        h3_counter = 0
        h4_counter = 0

        for i in range(s, e):
            p = paragraphs[i]
            level = get_heading_level(p)

            if level == 2:
                h2_counter += 1
                h3_counter = 0
                h4_counter = 0
                new_prefix = f"{ch_num}.{h2_counter}"
            elif level == 3:
                h3_counter += 1
                h4_counter = 0
                new_prefix = f"{ch_num}.{h2_counter}.{h3_counter}"
            elif level == 4:
                h4_counter += 1
                new_prefix = f"{ch_num}.{h2_counter}.{h3_counter}.{h4_counter}"
            else:
                continue

            old_text = p.text.strip()
            # Match existing number prefix (e.g., "1.2.3 " or "5.1 ")
            new_text = re.sub(
                r"^\d+\.\d+(\.\d+)*\s*",
                f"{new_prefix} ",
                old_text,
            )
            if new_text != old_text:
                for run in p.runs:
                    run.text = re.sub(
                        r"^\d+\.\d+(\.\d+)*\s*",
                        f"{new_prefix} ",
                        run.text,
                    )
                changes += 1

    print(f"Modified {changes} headings")
    doc.save(doc_path)
    print(f"Saved: {doc_path}")
